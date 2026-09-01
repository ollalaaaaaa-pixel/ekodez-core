import hashlib
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


class DocumentTemplateError(ValueError):
    pass


@dataclass(frozen=True)
class PackageFile:
    kind: str
    name: str
    size: int
    sha256: str


@dataclass(frozen=True)
class PackageManifest:
    directory: Path
    version: int
    files: tuple[PackageFile, ...]


TEMPLATES = {
    "inspection": ("inspection-act.docx", "Акт_осмотра.docx"),
    "work_act": ("work-completion-act.docx", "Акт_выполненных_работ.docx"),
    "invoice": ("invoice.docx", "Счёт.docx"),
}
TOKEN_PATTERN = re.compile(r"\{\{([^{}]+)\}\}")
INVALID_FILENAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _safe_component(value: str) -> str:
    cleaned = INVALID_FILENAME.sub("_", value).strip(" .")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return (cleaned or "Объект")[:80]


def _tokens_in_package(path: Path) -> set[str]:
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", "ignore")
            for name in archive.namelist()
            if name.endswith(".xml")
        )
    return set(TOKEN_PATTERN.findall(xml))


def _replace_in_paragraph(paragraph: Paragraph, values: dict[str, str]) -> None:
    while True:
        combined = "".join(run.text for run in paragraph.runs)
        match = TOKEN_PATTERN.search(combined)
        if match is None:
            return
        token = match.group(1)
        replacement = values[token]
        start, end = match.span()
        position = 0
        start_index = end_index = 0
        start_offset = end_offset = 0
        for index, run in enumerate(paragraph.runs):
            next_position = position + len(run.text)
            if position <= start < next_position:
                start_index = index
                start_offset = start - position
            if position < end <= next_position:
                end_index = index
                end_offset = end - position
                break
            position = next_position
        start_run = paragraph.runs[start_index]
        if start_index == end_index:
            start_run.text = (
                start_run.text[:start_offset]
                + replacement
                + start_run.text[end_offset:]
            )
            continue
        suffix = paragraph.runs[end_index].text[end_offset:]
        start_run.text = start_run.text[:start_offset] + replacement
        for index in range(start_index + 1, end_index):
            paragraph.runs[index].text = ""
        paragraph.runs[end_index].text = suffix


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell):
    yield from cell.paragraphs
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _iter_paragraphs(document: DocumentType):
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _fill_template(source: Path, destination: Path, values: dict[str, str]) -> None:
    document = Document(str(source))
    for paragraph in _iter_paragraphs(document):
        _replace_in_paragraph(paragraph, values)
    document.save(str(destination))
    remaining = _tokens_in_package(destination)
    if remaining:
        raise DocumentTemplateError(
            "unfilled template fields: " + ", ".join(sorted(remaining))
        )


def _next_version(parent: Path) -> int:
    versions = [
        int(match.group(1))
        for path in parent.glob("v*")
        if (match := re.fullmatch(r"v(\d+)", path.name)) is not None
    ]
    return max(versions, default=0) + 1


def resolve_package_file(
    *,
    output_root: Path,
    object_name: str,
    period_month: date,
    version: int,
    file_name: str,
) -> Path:
    root = output_root.resolve()
    path = (
        root
        / period_month.strftime("%Y-%m")
        / _safe_component(object_name)
        / f"v{version}"
        / file_name
    ).resolve()
    if not path.is_relative_to(root):
        raise DocumentTemplateError("invalid document path")
    return path


def build_month_package(
    *,
    template_dir: Path,
    output_root: Path,
    object_name: str,
    period_month: date,
    paid_service_due: bool,
    values: dict[str, str],
) -> PackageManifest:
    selected = ["inspection"]
    if paid_service_due:
        selected.extend(("work_act", "invoice"))

    sources = {kind: template_dir / TEMPLATES[kind][0] for kind in selected}
    missing_files = [str(path) for path in sources.values() if not path.is_file()]
    if missing_files:
        raise DocumentTemplateError("missing templates: " + ", ".join(missing_files))
    required = set().union(*(_tokens_in_package(path) for path in sources.values()))
    missing_values = sorted(required - values.keys())
    if missing_values:
        raise DocumentTemplateError(
            "missing template values: " + ", ".join(missing_values)
        )

    root = output_root.resolve()
    parent = root / period_month.strftime("%Y-%m") / _safe_component(object_name)
    parent.mkdir(parents=True, exist_ok=True)
    version = _next_version(parent)
    final_directory = parent / f"v{version}"
    temporary = Path(tempfile.mkdtemp(prefix=".package-", dir=parent))
    try:
        files: list[PackageFile] = []
        for kind in selected:
            destination = temporary / TEMPLATES[kind][1]
            _fill_template(sources[kind], destination, values)
            content = destination.read_bytes()
            files.append(
                PackageFile(
                    kind=kind,
                    name=destination.name,
                    size=len(content),
                    sha256=hashlib.sha256(content).hexdigest(),
                )
            )
        temporary.rename(final_directory)
    except Exception:
        shutil.rmtree(temporary, ignore_errors=True)
        raise
    return PackageManifest(
        directory=final_directory,
        version=version,
        files=tuple(files),
    )
